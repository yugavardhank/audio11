from fpdf import FPDF
from docx import Document

def export_pdf(result, path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=10)

    pdf.cell(0, 10, "Audio Analysis Report", ln=True)

    topics = (result or {}).get("topics") or []
    metrics = (result or {}).get("metrics") or {}
    accuracy = metrics.get("accuracy", "N/A")

    pdf.cell(0, 10, f"Topics: {len(topics)}", ln=True)
    pdf.cell(0, 10, f"Accuracy score: {accuracy}", ln=True)

    for t in topics:
        title = (t or {}).get("title") if isinstance(t, dict) else getattr(t, "title", "")
        summary = (t or {}).get("summary") if isinstance(t, dict) else getattr(t, "summary", "")
        pdf.multi_cell(0, 8, f"{title}\n{summary}\n")

    pdf.output(path)

def export_docx(result, path):
    doc = Document()
    doc.add_heading("Audio Analysis Report", 1)

    topics = (result or {}).get("topics") or []
    for t in topics:
        title = (t or {}).get("title") if isinstance(t, dict) else getattr(t, "title", "")
        summary = (t or {}).get("summary") if isinstance(t, dict) else getattr(t, "summary", "")
        doc.add_heading(title or "Untitled Topic", 2)
        if summary:
            doc.add_paragraph(summary)

    doc.save(path)
